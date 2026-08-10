"""
Written-warning document generation.

Takes a graded call and produces a .docx warning: a standard HR header block
(warning level, purpose, signatures) and one narrative finding per section of
**that organization's own checklist** that the call failed, each citing a
transcript timestamp.

This file was inherited from the single-tenant tool it grew out of, and until
recently still had that one customer baked into every document it produced —
their name in the system prompt, their sales process described to the model as
context, their logo path, and a hardcoded job title and department asserted
about whoever the warning was for. Nothing here is specific to any tenant now,
and nothing should become so again: everything that varies is a parameter.
"""

import io
import json
import logging

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches

# One model choice for the whole app. This file used to pin its own, and had
# drifted a full major version behind the grader whose findings it writes up.
from pipeline import MODEL

logger = logging.getLogger('qaboom.writeup')


# ---------------------------------------------------------------------------
# Claude call to produce the two prose finding bodies
# ---------------------------------------------------------------------------

_WRITEUP_SYSTEM_PROMPT = """You write the findings section of an employee \
written warning for a call-centre quality review.

You will be given:
- the employing organization's name
- the agent's name
- the full call transcript, with [MM:SS] or [HH:MM:SS] timestamps
- the requirements this call failed, exactly as that organization wrote them
- any prohibited phrases detected, with the sentence that triggered them

Write ONE short finding for each failed requirement group you are given.

Rules:
- 1-3 sentences each. Formal, factual, HR/compliance tone.
- Start with the agent's literal name. Not "the agent", no placeholders.
- Cite at least one timestamp inline as [MM:SS], taken from the material you
  were given. Never invent one. If no timestamp supports the finding, say the
  requirement was not addressed at any point in the call, with no timestamp.
- Describe ONLY what this organization's own requirements say. You do not know
  their business, their products, their sales process, or what their agents are
  normally expected to say. Do not infer a workflow and do not import
  assumptions from any other company.
- Never assert a failure that is not in the material you were given. This
  document goes in someone's employment file.
- No bullet lists, no long quotes inside a finding.

Output STRICT JSON, no prose around it:
{"findings": [{"heading": "<the requirement group, as given>",
               "body": "<your 1-3 sentences>"}]}
"""


def failed_groups(report_json: dict | None) -> list[dict]:
    """The sections this call actually failed, with their missed requirements.

    Drives the document. The predecessor to this file hardcoded two headings
    from one company's warning template — every org got "Failure to Accurately
    Explain Program" whether or not that was anything to do with their
    checklist, or their business.
    """
    groups = []
    for section in (report_json or {}).get("sections") or []:
        if not isinstance(section, dict):
            continue
        missed = [
            i["name"] for i in section.get("items") or []
            if isinstance(i, dict) and i.get("name")
            and i.get("required", True) and i.get("status") != "covered"
        ]
        if missed:
            groups.append({"heading": section.get("name") or "Requirements",
                           "missed": missed})
    return groups


def prohibited_hits(report_json: dict | None) -> list[dict]:
    af = (report_json or {}).get("auto_fail_phrases") or {}
    return [p for p in (af.get("phrases") or []) if isinstance(p, dict)]


def generate_finding_bodies(claude_client, org_name, agent_name, transcript,
                            report_json, model=MODEL):
    """One narrative finding per failed section. Returns [] when nothing failed.

    Returning an empty list matters: a call with no failures must produce no
    findings rather than prose inventing some, because this text ends up in an
    employment record.
    """
    groups = failed_groups(report_json)
    hits = prohibited_hits(report_json)
    if not groups and not hits:
        return []

    material = {
        "organization": org_name,
        "agent": agent_name,
        "failed_requirements": groups,
        "prohibited_phrases": [
            {"phrase": h.get("phrase", ""), "timestamp": h.get("timestamp"),
             "quote": h.get("quote", "")}
            for h in hits
        ],
    }

    resp = claude_client.messages.create(
        model=model,
        max_tokens=1500,
        thinking={"type": "disabled"},
        system=_WRITEUP_SYSTEM_PROMPT,
        messages=[{"role": "user", "content":
                   f"{json.dumps(material, indent=2)}\n\nTRANSCRIPT:\n{transcript}"}],
    )

    raw = resp.content[0].text
    i, j = raw.find("{"), raw.rfind("}") + 1
    if i == -1 or j <= i:
        raise ValueError(f"No JSON in the model response: {raw[:200]!r}")
    findings = json.loads(raw[i:j]).get("findings") or []

    # Keep only headings we actually asked about. A heading the model invented
    # would be an unsupported allegation in an HR document.
    allowed = {g["heading"] for g in groups} | {"Prohibited language"}
    return [f for f in findings
            if isinstance(f, dict) and f.get("body")
            and f.get("heading") in allowed]


# ---------------------------------------------------------------------------
# .docx construction
#
# Findings are one bullet per failed section of the org's own checklist, so a
# document only ever asserts what that org actually requires. The predecessor
# emitted two fixed headings from one company's warning template.
# ---------------------------------------------------------------------------

def _set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _add_runs(paragraph, segments):
    """segments: list of (text, bold) tuples."""
    for text, bold in segments:
        run = paragraph.add_run(text)
        run.bold = bold


def _set_paragraph_spacing(paragraph, space_before_pt=0, space_after_pt=4, line_spacing=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing = line_spacing


def _build_header_table(doc, agent_name, mode='written'):
    # Job title, department, hire date and manager are left blank for whoever
    # signs the document. The app does not hold any of them, and a warning
    # that states an employee's job title incorrectly is worse than one that
    # leaves a line to fill in.
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.autofit = False

    row0 = table.rows[0].cells
    row0[1].merge(row0[2])
    _add_runs(row0[0].paragraphs[0], [('Employee Name: ', True), (agent_name, False)])
    _add_runs(row0[1].paragraphs[0], [('Job Title: ', True)])

    row1 = table.rows[1].cells
    row1[1].merge(row1[2])
    _add_runs(row1[0].paragraphs[0], [('Date of Hire: ', True)])
    _add_runs(row1[1].paragraphs[0], [('Department: ', True)])

    row2 = table.rows[2].cells
    row2[1].merge(row2[2])
    _add_runs(row2[0].paragraphs[0], [('Date of Counseling: ', True)])
    _add_runs(row2[1].paragraphs[0], [('Supervisor/Manager: ', True)])

    # Verbal warning pre-marks Oral Discussion + First Warning per the template.
    if mode == 'verbal':
        row3_labels = ('_x_ Oral Discussion', '_x_ First Warning', '____ Second Warning')
    else:
        row3_labels = ('___ Oral Discussion', '____ First Warning', '____ Second Warning')
    row3 = table.rows[3].cells
    for cell, label in zip(row3, row3_labels):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run(label).bold = True

    row4 = table.rows[4].cells
    for cell, label in zip(row4, ('____ Final Warning', '___ Termination Decision', '__ PIP')):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run(label).bold = True

    row5 = table.rows[5].cells
    row5[0].merge(row5[1]).merge(row5[2])
    purpose_noun = 'warning' if mode == 'verbal' else 'meeting'
    row5[0].paragraphs[0].add_run(f'The purpose of this {purpose_noun} is to address:')

    purpose_row = table.add_row().cells
    purpose_row[1].merge(purpose_row[2])
    purpose_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    purpose_row[0].paragraphs[0].add_run('___ Performance Conduct')
    purpose_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    purpose_row[1].paragraphs[0].add_run('____ Policy/Procedure Violation')

    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_org_heading(doc, org_name):
    """The employer's name, as plain text.

    This used to load a bundled logo file belonging to the single customer this
    tool was originally built for. That file is not in this repo, so what every
    generated document actually contained was the literal placeholder text.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(org_name or "")
    run.bold = True
    run.font.size = Pt(14)
    _set_paragraph_spacing(p, space_after_pt=10)


def _add_main_bullet(doc, heading, body_text):
    p = doc.add_paragraph(style='List Bullet')
    _add_runs(p, [(heading, True), (' ', True), (body_text, False)])
    _set_paragraph_spacing(p, space_after_pt=6)
    return p


def build_writeup_docx(org_name, agent_name, internal_id, call_date,
                       findings, mode='written'):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ---- Page 1 ----
    _add_org_heading(doc, org_name)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('DOCUMENTED VERBAL WARNING' if mode == 'verbal' else 'WRITTEN WARNING')
    run.bold = True
    run.font.size = Pt(14)

    _build_header_table(doc, agent_name, mode=mode)

    doc.add_paragraph()  # spacer

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Both the internal ID and the call date are optional on upload, so each
    # clause is omitted rather than printing "file # , dated ,".
    opening = [('Following a review of a client call', False)]
    if internal_id:
        opening += [(', file # ', False), (str(internal_id), False)]
    if call_date:
        opening += [(', dated ', False), (call_date, False)]
    opening += [
        (', a compliance review of this call identified the requirement failures '
         'set out below. Each is a departure from the standards this role is '
         'expected to meet on every client interaction.', False),
    ]
    _add_runs(intro, opening)
    _set_paragraph_spacing(intro, space_after_pt=8)

    p = doc.add_paragraph()
    p.add_run('Findings include:')
    _set_paragraph_spacing(p, space_after_pt=4)

    # One bullet per failed section of this org's checklist. The version of this
    # file inherited from the single-tenant tool asserted, in every document it
    # produced, that the agent "failed to properly disclose missed payments" —
    # true of one customer's business and of no other, and not necessarily true
    # of the call in hand either. An HR document may only allege what was found.
    for f in findings:
        heading = (f.get('heading') or '').strip().rstrip(':')
        _add_main_bullet(doc, f'{heading}:', f.get('body') or '')

    # Corrective Action follows the findings directly (no forced page break)
    # so short calls don't leave a large blank gap at the bottom of page 1;
    # content reflows onto page 2 only when it actually runs out of room.
    ca_heading = doc.add_paragraph()
    ca_heading.add_run('Corrective Action').bold = True
    _set_paragraph_spacing(ca_heading, space_before_pt=10, space_after_pt=4)
    ca_heading.paragraph_format.keep_with_next = True

    corrective_items = [
        'Follow all company compliance policies, procedures and applicable '
        'regulatory standards on every call.',
        'Complete each required step of the approved call checklist in full, in '
        'the manner the checklist specifies.',
        'Provide clients with accurate and complete disclosures, and confirm '
        'their understanding before proceeding.',
        'Cease any misleading or prohibited language during all client '
        'interactions.',
        'Seek guidance from management immediately when a situation is unclear.',
        'Attend one-on-one coaching with the manager to obtain a clear '
        'understanding of the job expectations.',
    ]
    for item in corrective_items:
        b = doc.add_paragraph(style='List Bullet')
        b.add_run(item)
        _set_paragraph_spacing(b, space_after_pt=3)

    # The closing monitoring paragraph, employee comments, signature lines and
    # disclaimer appear only on the WRITTEN warning. The documented verbal
    # warning ends after the Corrective Action items (per its template).
    if mode != 'verbal':
        closing = doc.add_paragraph()
        closing.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        closing.add_run(
            'Your progress on the above corrective action items requiring improvement will be closely '
            'monitored. Improvement must begin immediately and be maintained. If no improvement in any '
            'and all of the above areas is noticed at any time following this corrective write-up, '
            'further disciplinary action up to and including termination may occur.'
        )
        _set_paragraph_spacing(closing, space_after_pt=10)

        p = doc.add_paragraph()
        p.add_run('Employee’s Comments:').bold = True
        # Keep the underscore count within the text column so it stays a single
        # line — too many wraps to a short stub on the next line.
        doc.add_paragraph('_' * 80)

        sig_spacer = doc.add_paragraph()
        _set_paragraph_spacing(sig_spacer, space_after_pt=2)

        sig1 = doc.add_paragraph()
        _add_runs(sig1, [
            ('Employee Signature*', True),
            ('________________________   ', False),
            ('Date:', True),
            ('___________________', False),
        ])
        _set_paragraph_spacing(sig1, space_after_pt=8)

        sig2 = doc.add_paragraph()
        _add_runs(sig2, [
            ('Manager’s Signature ', True),
            ('________________________   ', False),
            ('Date:', True),
            ('___________________', False),
        ])
        _set_paragraph_spacing(sig2, space_after_pt=14)

        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = disclaimer.add_run(
            '*It is understood that the employee’s signature indicates the information has been '
            'discussed and does not necessarily Indicate agreement*'
        )
        r.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Top-level entry point used by the Flask route
# ---------------------------------------------------------------------------

