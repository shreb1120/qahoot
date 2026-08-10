"""Written-warning generation.

This file exists mostly because of what it replaced. `writeup.py` was inherited
from the single-tenant tool this product grew out of, and every document it
generated named that one customer, asserted a job title and department nobody
had entered, and alleged a specific violation ("failed to properly disclose
missed payments") whether or not that was among the requirements the call
actually failed — or even among the requirements the org has.

A written warning goes in someone's employment file. The tests below are about
the document never asserting anything that wasn't found.
"""
import json

import pytest

import writeup


class FakeMessages:
    """Stands in for the Anthropic client, and records what it was told."""

    def __init__(self, findings):
        self._findings = findings
        self.last_kwargs = None

    def create(self, **kwargs):
        self.last_kwargs = kwargs
        body = json.dumps({"findings": self._findings})
        return type("R", (), {"content": [type("C", (), {"text": body})()]})()


class FakeClient:
    def __init__(self, findings):
        self.messages = FakeMessages(findings)


REPORT = {
    "sections": [
        {"name": "Approval script", "items": [
            {"name": "Disclose fees", "status": "missed", "required": True},
            {"name": "Confirm identity", "status": "covered", "required": True},
        ]},
        {"name": "Post-enrollment", "items": [
            {"name": "Explain cancellation", "status": "covered", "required": True},
        ]},
    ],
    "auto_fail_phrases": {"phrases": []},
}


# ───────────────── what the document may assert ─────────────────

def test_only_failed_sections_become_findings():
    """A section the agent passed must not appear. The old template printed two
    fixed headings regardless of the report."""
    groups = writeup.failed_groups(REPORT)
    assert [g["heading"] for g in groups] == ["Approval script"]
    assert groups[0]["missed"] == ["Disclose fees"]


def test_a_clean_call_produces_no_findings_and_no_api_call():
    """There is nothing to warn about. Generating prose here would mean inventing
    an allegation, so it must not even ask the model."""
    client = FakeClient([{"heading": "Approval script", "body": "x"}])
    clean = {"sections": [{"name": "Approval script", "items": [
        {"name": "Disclose fees", "status": "covered", "required": True}]}]}

    assert writeup.generate_finding_bodies(client, "Acme", "Dana", "t", clean) == []
    assert client.messages.last_kwargs is None, "asked the model about a passing call"


def test_a_heading_the_model_invented_is_dropped():
    """The model returning a section that isn't in the report would put an
    unsupported allegation in an employment record."""
    client = FakeClient([
        {"heading": "Approval script", "body": "Dana did not disclose fees [02:11]."},
        {"heading": "Timeliness", "body": "Dana was late."},
    ])
    out = writeup.generate_finding_bodies(client, "Acme", "Dana", "t", REPORT)
    assert [f["heading"] for f in out] == ["Approval script"]


def test_optional_items_do_not_trigger_a_warning():
    report = {"sections": [{"name": "Rapport", "items": [
        {"name": "Use the client's name", "status": "missed", "required": False}]}]}
    assert writeup.failed_groups(report) == []


def test_a_malformed_report_does_not_crash_the_download():
    for bad in (None, {}, {"sections": None}, {"sections": [None, "x"]},
                {"sections": [{"items": [None, {"status": "missed"}]}]}):
        assert isinstance(writeup.failed_groups(bad), list)


# ───────────────── nothing tenant-specific ─────────────────

def test_the_prompt_names_no_company_and_assumes_no_business():
    """The regression that made this file a liability: one customer's name and
    sales process were in the system prompt, so every other org's warnings were
    written against a workflow that had nothing to do with them."""
    prompt = writeup._WRITEUP_SYSTEM_PROMPT.lower()
    for term in ("better debt", "bds", "alleviate", "debt advisor",
                 "enrollment", "loan decline", "escrow", "settlement"):
        assert term not in prompt, f"{term!r} is back in the write-up prompt"


def test_the_org_name_is_passed_to_the_model_not_hardcoded():
    client = FakeClient([{"heading": "Approval script", "body": "b"}])
    writeup.generate_finding_bodies(client, "Northwind Lending", "Dana", "t", REPORT)
    sent = client.messages.last_kwargs["messages"][0]["content"]
    assert "Northwind Lending" in sent


def test_the_writeup_model_tracks_the_grader():
    """These drifted a major version apart while nobody was looking."""
    import pipeline
    assert writeup.MODEL == pipeline.MODEL


# ───────────────── the document itself ─────────────────

def _docx_text(buf):
    from docx import Document
    doc = Document(buf)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        parts += [c.text for r in t.rows for c in r.cells]
    return "\n".join(parts)


def test_the_document_carries_the_orgs_own_name_and_findings():
    buf = writeup.build_writeup_docx(
        "Northwind Lending", "Dana Reyes", "CALL-9", "2026-08-01",
        [{"heading": "Approval script",
          "body": "Dana Reyes did not disclose fees [02:11]."}],
    )
    text = _docx_text(buf)
    assert "Northwind Lending" in text
    assert "Dana Reyes" in text
    assert "Approval script" in text
    assert "[02:11]" in text


def test_the_document_states_no_job_title_department_or_stray_allegation():
    """Every field here was previously filled in with one customer's values, for
    every employee of every organization."""
    buf = writeup.build_writeup_docx(
        "Northwind Lending", "Dana Reyes", "CALL-9", "2026-08-01",
        [{"heading": "Approval script", "body": "Dana Reyes did not disclose fees [02:11]."}],
    )
    text = _docx_text(buf).lower()
    for term in ("sr debt advisor", "bds", "better debt", "missed payments",
                 "post-enrollment script", "[bds logo]"):
        assert term not in text, f"{term!r} is back in the generated document"

    # The blank fields are still there to be filled in by hand.
    assert "job title:" in text and "department:" in text
